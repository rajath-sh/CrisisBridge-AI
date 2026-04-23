import os
from pathlib import Path
from dataclasses import dataclass
from typing import List

# Conditional imports for PDF/DOCX (if needed later)
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None


@dataclass
class Document:
    """A loaded document with text and metadata."""
    content: str
    metadata: dict


class DocumentLoader:
    """
    Loads text, PDF, and DOCX files from a directory.
    Extracts content and basic metadata.
    """

    def __init__(self, directory_path: str):
        self.directory_path = Path(directory_path)

    def load_all(self) -> List[Document]:
        """Loads all supported documents in the directory."""
        documents = []
        
        if not self.directory_path.exists():
            print(f"Directory {self.directory_path} does not exist.")
            return documents

        for file_path in self.directory_path.rglob("*"):
            if not file_path.is_file():
                continue
                
            doc = self._load_file(file_path)
            if doc:
                documents.append(doc)
                
        return documents

    def _load_file(self, file_path: Path) -> Document | None:
        """Dispatches to the correct loader based on file extension."""
        ext = file_path.suffix.lower()
        
        metadata = {"source": file_path.name, "path": str(file_path)}
        
        try:
            if ext == ".txt":
                content = file_path.read_text(encoding="utf-8")
                return Document(content=content, metadata=metadata)
                
            elif ext == ".pdf":
                if PdfReader is None:
                    print(f"Skipping {file_path}: pypdf not installed")
                    return None
                
                reader = PdfReader(str(file_path))
                text_parts = [page.extract_text() for page in reader.pages if page.extract_text()]
                content = "\n\n".join(text_parts)
                return Document(content=content, metadata=metadata)
                
            elif ext == ".docx":
                if docx is None:
                    print(f"Skipping {file_path}: python-docx not installed")
                    return None
                    
                doc = docx.Document(str(file_path))
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return Document(content=content, metadata=metadata)
                
        except Exception as e:
            print(f"Error loading {file_path}: {e}")
            
        return None
